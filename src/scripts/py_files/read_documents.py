"""Read and process aviation-related documents.

This module provides utilities to:

- extract text from PDFs (via pdfplumber and PyPDF2) and DOCX files
- clean and normalize raw text, including abbreviation expansion and word splitting
- perform NLP tasks such as entity recognition, POS tagging, and lemmatization
- extract document metadata
- classify documents with simple keyword-based heuristics
- checkpoint long-running processing so it can be safely resumed

The implementation is designed to be robust to noisy OCR and heterogeneous
document formats, with structured logging for observability and debugging.
"""

import csv
import logging
import os
import pickle
import re
import sys
import tempfile
import warnings
import docx
import nltk
import pdfplumber
import PyPDF2
import spacy
import wordninja
from config import (ABBREVIATIONS_CSV, DOCUMENTS_DIR, LOG_DIR, PKL_FILENAME,
                    TEXT_EXPANDED_DIR, TEXT_OUTPUT_DIR)
from docx import Document
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from spellchecker import SpellChecker

# Load spaCy's English model
nlp = spacy.load("en_core_web_sm")
nlp.max_length = 2000000  # or any other suitable value

# Initialize spellchecker
spell = SpellChecker()

# Suppress specific warnings
warnings.filterwarnings("ignore", message="usetex mode requires TeX.")

# Global stopwords
STOP_WORDS = set(stopwords.words("english"))

log_file_path = LOG_DIR / "read_documents.log"
LOG_LEVEL = os.getenv("READ_DOC_LOG_LEVEL", "INFO").upper()
CHECKPOINT_EVERY = max(1, int(os.getenv("READ_DOC_CHECKPOINT_EVERY", "1")))
NLP_CHUNK_CHARS = max(50000, int(os.getenv("READ_DOC_NLP_CHUNK_CHARS", "180000")))

# Ensure the log directory exists
os.makedirs(LOG_DIR, exist_ok=True)

# Configure logging properly
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(log_file_path, mode="w"),  # Ensure file is written
        logging.StreamHandler(),  # Print logs to console as well
    ],
)

logging.info("Logging initialized successfully.")
for noisy_logger_name in ("pdfminer", "pdfplumber", "PIL", "urllib3"):
    logging.getLogger(noisy_logger_name).setLevel(logging.WARNING)

# Ensure directories exist
for directory in [TEXT_OUTPUT_DIR, TEXT_EXPANDED_DIR, PKL_FILENAME.parent]:
    os.makedirs(directory, exist_ok=True)  # No need to log this for every run


# Download NLTK data
def download_nltk_data():
    """Downloads required NLTK data only if not already installed.
    Args: None
    Returns: None
    """
    try:
        nltk.data.find("tokenizers/punkt")
        nltk.data.find("corpora/stopwords")
        nltk.data.find("corpora/wordnet")
    except LookupError:
        nltk.download("punkt")
        nltk.download("stopwords")
        nltk.download("wordnet")


download_nltk_data()

# Create custom pipeline component for aviation NER
@spacy.Language.component("aviation_ner")
def aviation_ner(doc):
    """
    This custom spaCy pipeline component identifies aviation-specific entities such
    as aircraft models, airport codes, flight numbers, airlines, and aviation organizations.
    It uses regex patterns to find these entities in the text and adds them to the doc.ents
    list while ensuring no overlap with existing entities. The component is designed to
    enhance the NER capabilities of the spaCy model for aviation-related documents.
    Args:
    - doc (spacy.tokens.Doc): The spaCy Doc object to process.
    Returns:
    - spacy.tokens.Doc: The processed Doc object with added aviation entities.    
    """
    logging.info("Starting aviation_ner for document: %s...", doc[:50])
    patterns = [
        ("AIRCRAFT_MODEL", r"\b[A-Z]-?[0-9]{1,4}\b"),
        ("AIRPORT_CODE", r"\b[A-Z]{3}\b"),
        ("FLIGHT_NUMBER", r"\b[A-Z]{2,3}\s?[0-9]{1,4}\b"),
        (
            "AIRLINE",
            r"\b(American Airlines|Delta Air Lines|United Airlines|Southwest Airlines|Air France|Lufthansa|British Airways)\b",
        ),
        ("AVIATION_ORG", r"\b(FAA|EASA|ICAO|IATA)\b"),
    ]

    new_ents = []
    for ent_type, pattern in patterns:
        for match in re.finditer(pattern, doc.text):
            start, end = match.span()
            span = doc.char_span(start, end, label=ent_type)
            if span is not None:
                # Check for overlap with existing entities
                if not any(
                    span.start < ent.end and span.end > ent.start
                    for ent in list(doc.ents) + new_ents
                ):
                    new_ents.append(span)
                    logging.debug("Added new entity: %s (%s)", span.text, ent_type)

    doc.ents = list(doc.ents) + new_ents
    logging.info("Finished aviation_ner. Added %d new entities.", len(new_ents))
    return doc


# Add the custom component to the pipeline
nlp.add_pipe("aviation_ner", after="ner")


def load_abbreviation_dict():
    """
    Loads the abbreviation dictionary from a CSV file.
    Returns:
        dict: A dictionary mapping abbreviations to their full forms.
    """
    abbreviation_dict = {}
    try:
        for encoding in ("utf-8", "cp1252", "latin-1"):
            try:
                with open(ABBREVIATIONS_CSV, mode="r", encoding=encoding) as infile:
                    reader = csv.reader(infile)
                    abbreviation_dict = {
                        rows[0].strip(): rows[1].strip()
                        for rows in reader
                        if len(rows) >= 2
                    }
                break
            except UnicodeDecodeError:
                continue
    except FileNotFoundError:
        logging.error("Error: The file '%s' was not found.", ABBREVIATIONS_CSV)
    except (OSError, csv.Error) as err:
        logging.exception(
            "An error occurred while loading the abbreviation dictionary: %s", err
        )
    return abbreviation_dict


def split_connected_words_improved(text):
    """
    Split connected words while preserving important patterns
    like section references (e.g., "4.1.2", "§ 5.1"), aviation
    model numbers (e.g., "A320-214"), and common phrases
    (e.g., "Level 1", "Part 91").
    Args:    text (str): The input text to process.
    Returns:     str: The processed text with connected words split appropriately.    
    """
    words = re.findall(r"\w+|\W+", text)
    split_words = []

    for word in words:
        # ✅ Keep section references like "4.1.2", "§ 5.1" intact
        if re.match(r"^\d+(\.\d+)+$", word) or re.match(r"^§\s*\d+(\.\d+)*$", word):
            split_words.append(word)
            continue

        # ✅ Keep "Level 1", "Part 91"
        if re.match(r"^[A-Za-z]+\s\d+$", word):
            split_words.append(word)
            continue

        # ✅ Keep aviation model numbers (e.g., "A320-214") intact
        if re.match(r"^[A-Za-z]+\d+[-]?\d*$", word):
            split_words.append(word)
            continue

        # ✅ Process long alphanumeric words but keep numbers
        if len(word) > 15 and word.isalnum():
            split_parts = re.findall("[A-Z][a-z]*|[a-z]+|[0-9]+", word)
            split_words.extend(split_parts)
        else:
            split_words.append(word)

    # Apply wordninja ONLY to words, NOT numbers
    processed_text = " ".join(split_words)
    processed_text = " ".join(wordninja.split(processed_text))

    return processed_text


def filter_non_sense_strings(text):
    """
    Filter out non-sense strings that are likely OCR errors or artifacts, while preserving
    important alphanumeric patterns and section references.
    Args:    text (str): The input text to filter.
    Returns:     str: The filtered text with non-sense strings removed.    
    """
    words = text.split()
    cleaned_words = []
    for word in words:
        # Keep numbers and section references like "5.1.2" or "Part 5"
        if re.match(r"^[a-zA-Z0-9.\-]+$", word) and len(set(word.lower())) > 2:
            cleaned_words.append(word)
    return " ".join(cleaned_words)


def extract_section_references(text):
    """Extract structured numerical references like '4.1.2' and 'Level 1'.
    Args:      text (str): The input text from which to extract references.
    Returns:     list: A list of extracted section references found in the text.
    """

    # Updated regex patterns
    section_pattern = re.findall(r"\b\d+(\.\d+)+\b", text)  # Matches "4.1.2"
    level_pattern = re.findall(r"\b[Ll]evel\s\d+\b", text)  # Matches "Level 1"
    part_pattern = re.findall(r"\b[Pp]art\s\d+\b", text)  # Matches "Part 121"
    paragraph_pattern = re.findall(r"^§\s*\d+(\.\d+)*$", text)  # Matches "§ 5.1"

    references = section_pattern + level_pattern + part_pattern + paragraph_pattern

    # Debugging output
    if references:
        logging.info("Extracted Section References: %s", references)
    else:
        logging.warning("No section references found in text.")

    return references if references else ["No References Found"]


def preprocess_text_with_sentences(text):
    """
    Preprocess the input text by tokenizing it into sentences, lemmatizing the words,
    and filtering out stop words and non-alphanumeric tokens while preserving important patterns.
    Args:    text (str): The input text to preprocess.
    Returns:    
    str: The preprocessed text with sentences preserved and cleaned.    
    """
    sentences = []

    for doc in iter_nlp_docs(text):
        for sent in doc.sents:
            cleaned_sentence = " ".join(
                token.lemma_.lower()
                for token in sent
                if (
                    token.is_alpha
                    or token.like_num
                    or re.match(r"^\d+(\.\d+)*$", token.text)
                )  # Preserve numbers
                and token.text.lower() not in STOP_WORDS
            )
            if cleaned_sentence:
                sentences.append(cleaned_sentence)

    return " ".join(sentences)


def extract_personal_names(text):
    """
    Extract personal names from the text using spaCy's NER capabilities.
    Args:    text (str): The input text from which to extract personal names.
    Returns:     list: A list of personal names found in the text.    
    """
    names = []
    for doc in iter_nlp_docs(text):
        names.extend([ent.text for ent in doc.ents if ent.label_ == "PERSON"])
    return names


def extract_entities_and_pos_tags(text):
    """
    Extract named entities and part-of-speech tags from the text using spaCy's NLP capabilities.
    This function processes the input text in manageable chunks to avoid exceeding the NLP model's maximum length.
    It collects entities and POS tags while ensuring that important patterns like section references are also captured as entities.
    Args:    text (str): The input text from which to extract entities and POS tags.
    Returns:     tuple: A tuple containing a list of entities (text and label) and a list of POS tags (text and tag).
    """
    entities = []
    pos_tags = []
    for doc in iter_nlp_docs(text):
        entities.extend([(ent.text, ent.label_) for ent in doc.ents])
        pos_tags.extend([(token.text, token.pos_) for token in doc])
    # References as entities
    section_references = extract_section_references(text)
    for ref in section_references:
        entities.append((ref, "SECTION_REF"))

    return entities, pos_tags


def split_text_for_nlp(text, max_chars=NLP_CHUNK_CHARS):
    """
    Split the input text into chunks that are suitable for processing
    by the NLP model, while trying to preserve sentence boundaries and 
    important patterns.
    Args:    text (str): The input text to split.
    max_chars (int): The maximum number of characters for each chunk.
    Returns:     list: A list of text chunks that are suitable for NLP processing.
    """
    if not text:
        return []

    if len(text) <= max_chars:
        return [text]

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + max_chars, text_len)
        if end < text_len:
            split_idx = max(
                text.rfind("\n", start, end),
                text.rfind(". ", start, end),
                text.rfind(" ", start, end),
            )
            if split_idx > start + int(max_chars * 0.5):
                end = split_idx + 1

        chunks.append(text[start:end])
        start = end

    return chunks


def iter_nlp_docs(text):
    """
    Iterate over chunks of text and process them with the NLP model.
    This function splits the input text into manageable chunks to avoid exceeding
    the NLP model's maximum length and yields processed Doc objects.
    Args:
    - text (str): The input text to process.
    Yields:
    - spacy.tokens.Doc: Processed Doc objects for each chunk of text.
    Returns: None    
    """
    safe_limit = max(1000, nlp.max_length - 1000)
    chunk_size = min(NLP_CHUNK_CHARS, safe_limit)

    for text_chunk in split_text_for_nlp(text, max_chars=chunk_size):
        cleaned_chunk = text_chunk.strip()
        if not cleaned_chunk:
            continue
        yield nlp(cleaned_chunk)


def expand_abbreviations_in_text(text, abbreviation_dict):
    """
    Expand abbreviations in the input text using the provided abbreviation dictionary.
    Args:    text (str): The input text containing potential abbreviations.
    abbreviation_dict (dict): A dictionary mapping abbreviations to their full forms.
    Returns:     str: The text with abbreviations expanded to their full forms.   
    """
    words = text.split()
    expanded_words = []
    for word in words:
        if word.lower() in abbreviation_dict:
            expanded_words.append(abbreviation_dict[word.lower()])
        else:
            expanded_words.append(word)
    return " ".join(expanded_words)


def extract_text_from_pdf_with_pdfplumber(pdf_path):
    """
    Extract text from PDF using pdfplumber.

    Args:
        pdf_path (str): Path to the PDF file.

    Returns:
        str: Extracted text from the PDF.
    """
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = "".join([(page.extract_text() or "") + "\n" for page in pdf.pages])
            return text
    except Exception as e:
        print(f"Failed to process PDF {pdf_path}: {e}")
        return ""


def extract_text_from_pdf_with_pypdf2(pdf_path):
    """Extract text from PDF using PyPDF2 as an alternative parser.
    Args:    pdf_path (str): Path to the PDF file.
    Returns:     str: Extracted text from the PDF, or an empty string if extraction fails.
    """
    try:
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            return "\n".join([(page.extract_text() or "") for page in reader.pages])
    except Exception as e:
        logging.warning("PyPDF2 extraction failed for %s: %s", pdf_path, e)
        return ""


def text_quality_score(text):
    """
    Score extracted text quality using length and alphanumeric ratio.
    Args:    text (str): The input text to score.
    Returns:     float: Quality score between 0.0 and 1.0.
    """
    if not text:
        return 0.0
    stripped = text.strip()
    if not stripped:
        return 0.0

    alnum_count = sum(ch.isalnum() for ch in stripped)
    alnum_ratio = alnum_count / max(len(stripped), 1)
    length_score = min(len(stripped) / 2000, 1.0)  # Normalize around 2k chars
    return (0.7 * length_score) + (0.3 * alnum_ratio)


def extract_text_from_pdf(pdf_path):
    """
    Try multiple PDF extraction strategies and keep the best result.

    Returns:
        tuple[str, str, float]: (text, extraction_method, quality_score)
    """
    candidates = []

    text_pypdf2 = extract_text_from_pdf_with_pypdf2(pdf_path)
    candidates.append(("pypdf2", text_pypdf2, text_quality_score(text_pypdf2)))

    text_pdfplumber = extract_text_from_pdf_with_pdfplumber(pdf_path)
    candidates.append(
        ("pdfplumber", text_pdfplumber, text_quality_score(text_pdfplumber))
    )

    method, text, score = max(candidates, key=lambda item: item[2])
    return text, method, score


def extract_keywords(documents, top_n=10):
    """
    Extract keywords from the documents using TF-IDF vectorization.
    Args:    documents (list): A list of document dictionaries, each containing a 'text' field.
    top_n (int): The number of top keywords to extract for each document.
    Returns: None (the function modifies the documents in place by adding a 'keywords' field).
    """
    texts = [doc["text"] for doc in documents]

    # Debug: Check if we have any texts
    logging.debug(f"Number of documents: {len(texts)}")
    if len(texts) == 0:
        logging.error("No texts found in documents")
        return

    # Debug: Check the content of the first few texts
    for i, text in enumerate(texts[:5]):
        logging.debug(f"Text {i} (first 100 chars): {text[:100]}")
        logging.debug(f"Text {i} length: {len(text)}")

    vectorizer = TfidfVectorizer(stop_words="english", max_features=1000)

    try:
        tfidf_matrix = vectorizer.fit_transform(texts)
    except ValueError as e:
        logging.error(f"ValueError in TfidfVectorizer: {e}")
        logging.error("Vocabulary: " + str(vectorizer.vocabulary_))
        return


def extract_metadata(file_path):
    """
    Extract metadata from PDF or DOCX files.

    Args:
        file_path (str): Path to the file.

    Returns:
        dict: Metadata extracted from the file.
    """
    metadata = {"title": "", "author(s)": "", "category": "", "tags": ""}

    # Handle PDF files
    if file_path.endswith(".pdf"):
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            pdf_metadata = reader.metadata or {}
            metadata["title"] = pdf_metadata.get("/Title", "")
            metadata["author(s)"] = pdf_metadata.get("/Author", "")

    elif file_path.endswith(".docx"):
        try:
            doc = docx.Document(file_path)
            # Look at the first few paragraphs for metadata
            text = "\n".join([p.text for p in doc.paragraphs[:5]])

            # Use regex to find metadata
            title_match = re.search(r"Title:\s*(.*)", text)
            if title_match:
                metadata["title"] = title_match.group(1).strip()

            author_match = re.search(r"Author\(s\):\s*(.*)", text)
            if author_match:
                metadata["author(s)"] = author_match.group(1).strip()

            category_match = re.search(r"Category:\s*(.*)", text)
            if category_match:
                metadata["category"] = category_match.group(1).strip()

            tags_match = re.search(r"Tags:\s*(.*)", text)
            if tags_match:
                metadata["tags"] = tags_match.group(1).strip()

        except Exception as e:
            logging.info(f"Error reading DOCX metadata: {e}")

    return metadata


def persist_documents_snapshot(documents, output_path=PKL_FILENAME):
    """Persist progress atomically so long runs can resume safely after interruption.
    Args:    documents (list): The list of document dictionaries to persist.
    output_path (Path): The path where the documents should be saved as a pickle file.
    Returns: None (the function saves the documents to the specified path).    
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        delete=False, dir=output_path.parent, suffix=".tmp"
    ) as tmp_file:
        pickle.dump(documents, tmp_file)
        tmp_path = tmp_file.name
    os.replace(tmp_path, output_path)


def classify_document(text):
    """
    Classify a document based on its content using predefined keywords.
    This function uses a simple keyword matching approach to classify the document into
    categories such as 'safety', 'maintenance', 'operations', etc. It counts the occurrences
    of relevant keywords for each category and assigns the category with the highest score to the document.
    Args:   text (str): The input text of the document to classify.
    Returns:     str: The category assigned to the document based on keyword matching.    
    """
    keywords = {
        "safety": [
            "safety",
            "hazard",
            "risk",
            "incident",
            "accident",
            "system",
            "emergency",
            "prevention",
            "safety management",
            "safety culture",
            "safety audit",
            "safety inspection",
            "safety compliance",
            "safety training",
        ],
        "maintenance": [
            "maintenance",
            "repair",
            "overhaul",
            "inspection",
            "servicing",
            "maintenance schedule",
            "maintenance manual",
            "maintenance log",
            "maintenance record",
            "maintenance check",
            "maintenance procedure",
            "maintenance crew",
            "maintenance facility",
        ],
        "operations": [
            "flight",
            "takeoff",
            "landing",
            "crew",
            "pilot",
            "aircraft",
            "airplane",
            "operations manual",
            "flight operations",
            "flight plan",
            "flight schedule",
            "flight crew",
            "flight deck",
            "flight control",
            "air traffic control",
            "ground operations",
        ],
        "regulations": [
            "regulation",
            "compliance",
            "standard",
            "rule",
            "law",
            "FAA",
            "ICAO",
            "EASA",
            "aviation regulation",
            "aviation law",
            "aviation compliance",
            "aviation standard",
            "aviation rule",
            "regulatory requirement",
            "regulatory compliance",
        ],
        "quality": [
            "quality",
            "performance",
            "service",
            "customer",
            "satisfaction",
            "design",
            "quality assurance",
            "quality control",
            "quality management",
            "quality audit",
            "quality inspection",
            "quality standard",
            "quality improvement",
            "quality system",
            "quality policy",
        ],
        "accident_reports": [
            "accident",
            "crash",
            "collision",
            "investigation",
            "report",
            "NTSB",
            "FAA",
            "safety board",
            "incident",
            "fatality",
            "injury",
            "damage",
            "wreckage",
            "black box",
            "flight recorder",
            "accident analysis",
            "accident summary",
            "accident investigation",
            "accident report",
            "accident findings",
            "accident cause",
            "accident prevention",
        ],
        "training_education": [
            "training",
            "education",
            "certification",
            "course",
            "syllabus",
            "training program",
            "instructor",
            "trainee",
            "aviation school",
            "flight school",
            "training manual",
        ],
        "weather_environment": [
            "weather",
            "environment",
            "climate",
            "turbulence",
            "storm",
            "wind",
            "visibility",
            "meteorology",
            "weather report",
            "weather forecast",
            "environmental impact",
        ],
        "technology_innovation": [
            "technology",
            "innovation",
            "avionics",
            "automation",
            "AI",
            "artificial intelligence",
            "drone",
            "UAV",
            "unmanned aerial vehicle",
            "new technology",
            "technological advancement",
        ],
        "security": [
            "security",
            "threat",
            "terrorism",
            "hijacking",
            "security measures",
            "airport security",
            "security protocol",
            "security breach",
            "security incident",
            "cybersecurity",
        ],
        "finance_economics": [
            "finance",
            "economics",
            "cost",
            "budget",
            "funding",
            "investment",
            "economic impact",
            "financial report",
            "financial analysis",
            "revenue",
            "expense",
            "profit",
            "loss",
        ],
        "human_factors": [
            "human factors",
            "ergonomics",
            "fatigue",
            "stress",
            "workload",
            "human performance",
            "crew resource management",
            "CRM",
            "human error",
            "human-machine interaction",
        ],
        "emergency_response": [
            "emergency",
            "response",
            "rescue",
            "evacuation",
            "emergency procedures",
            "emergency landing",
            "emergency services",
            "first aid",
            "crisis management",
            "disaster response",
        ],
    }

    text_lower = text.lower()
    scores = {
        category: sum(1 for word in words if word in text_lower)
        for category, words in keywords.items()
    }
    return max(scores, key=scores.get)


def read_documents_from_directory(
    directory_path,
    text_output_dir=None,
    text_expanded_dir=None,
    existing_documents=None,
    checkpoint_callback=None,
    ):
    logging.info(f"Starting to read documents from {directory_path}")
    if existing_documents is None:
        existing_documents = []

    existing_files = {doc["filename"] for doc in existing_documents}
    new_documents = []
    abbreviation_dict = load_abbreviation_dict()
    lemmatizer = WordNetLemmatizer()

    for filename in os.listdir(directory_path):
        logging.info(f"Processing file: {filename}")
        if filename in existing_files:
            continue

        file_path = os.path.join(directory_path, filename)
        text = ""
        extraction_method = "docx"
        extraction_quality = 1.0
        extension = os.path.splitext(filename)[1].lower()

        if extension == ".pdf":
            logging.info(f"Extracting text from PDF: {filename}")
            text, extraction_method, extraction_quality = extract_text_from_pdf(
                file_path
            )
        elif extension == ".docx":
            logging.info(f"Extracting text from DOCX: {filename}")
            try:
                doc = Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as e:
                logging.error(f"Failed to process DOCX {filename}: {e}")
                print(f"Failed to process DOCX {filename}: {e}")
                continue
        else:
            logging.warning(f"Skipping unsupported file type: {filename}")
            continue

        logging.info(f"Preprocessing text from {filename}")
        if not text:
            logging.warning(f"No text extracted from {filename}")
            continue

        needs_review = False
        if extension == ".pdf" and extraction_quality < 0.20:
            needs_review = True
            logging.warning(
                "Low-quality PDF extraction for %s (method=%s, score=%.3f). "
                "Likely scanned/image PDF; consider OCR.",
                filename,
                extraction_method,
                extraction_quality,
            )

        logging.info(f"Processing {filename}...")

        expanded_text = expand_abbreviations_in_text(text, abbreviation_dict)

        raw_text = expanded_text

        # Debug BEFORE preprocessing
        logging.debug(f"Raw expanded text BEFORE processing: {expanded_text[:500]}")

        expanded_text = split_connected_words_improved(expanded_text)

        expanded_text = filter_non_sense_strings(expanded_text)

        # Debug AFTER preprocessing
        logging.debug(f"Expanded text AFTER preprocessing: {expanded_text[:500]}")

        section_references = extract_section_references(expanded_text)

        if not section_references:
            logging.warning(f"No section references found in {filename}")

        preprocessed_text = preprocess_text_with_sentences(expanded_text)

        personal_names = extract_personal_names(preprocessed_text)

        entities, pos_tags = extract_entities_and_pos_tags(preprocessed_text)

        tokens = []
        for doc_chunk in iter_nlp_docs(preprocessed_text):
            tokens.extend(
                [
                    token.text.lower()
                    for token in doc_chunk
                    if (
                        token.is_alpha
                        or token.like_num
                        or re.match(r"^\d+(\.\d+)*$", token.text)
                        or re.match(r"^[A-Za-z]+\s\d+$", token.text)
                        or re.match(r"^§\s*\d+(\.\d+)*$", token.text)
                    )
                ]
            )

        tokens_without_stopwords = [
            token for token in tokens if token not in STOP_WORDS
        ]

        lemmatized_tokens = [
            lemmatizer.lemmatize(token) for token in tokens_without_stopwords
        ]

        if text_expanded_dir:
            clean_filename = os.path.splitext(filename)[0]
            output_file_path = os.path.join(text_expanded_dir, f"{clean_filename}.txt")
            with open(output_file_path, "w", encoding="utf-8") as out_file:
                out_file.write(raw_text)
            logging.info(f"Expanded text saved to: {output_file_path}")
            logging.info(f"Expanded text saved to: {output_file_path}")

        if text_output_dir:
            clean_filename = os.path.splitext(filename)[0]
            output_file_path = os.path.join(text_output_dir, f"{clean_filename}.txt")
            logging.info(f"Processed text saved to: {output_file_path}")
            with open(output_file_path, "w", encoding="utf-8") as out_file:
                out_file.write(preprocessed_text)
            logging.info(f"Text saved to: {output_file_path}")

        logging.info(f"Finished processing all documents in {directory_path}")
        metadata = extract_metadata(file_path)
        document_category = classify_document(preprocessed_text)

        # Remove 'category' field if it already exists
        if "category" in metadata:
            del metadata["category"]

        metadata["source_type"] = "pdf" if extension == ".pdf" else "docx"
        metadata["extraction_method"] = extraction_method
        metadata["extraction_quality"] = round(extraction_quality, 3)
        metadata["needs_manual_review"] = needs_review

        new_documents.append(
            {
                "filename": filename,
                "text": preprocessed_text,
                "tokens": lemmatized_tokens,
                "section_references": section_references,  # ✅ Ensure this is included
                "personal_names": personal_names,
                "entities": entities,
                "pos_tags": pos_tags,
                "metadata": metadata,
                "category": document_category,
            }
        )

        # Debugging output
        logging.info(f"Stored Section References for {filename}: {section_references}")

        if checkpoint_callback and (len(new_documents) % CHECKPOINT_EVERY == 0):
            checkpoint_callback(existing_documents + new_documents)

    return existing_documents + new_documents


def main():
    """
    Function main: Orchestrates the document processing workflow. It checks for existing processed documents
    in a pickle file and either loads them or processes new documents from the specified directory.
    It applies keyword extraction, downloads necessary NLTK data, and saves the updated list
    of documents back to the pickle file. The function includes error handling
    to log exceptions and ensure that the script exits gracefully in case of errors.
    Args: None
    Returns: None (the function performs processing and saves results but does not return a value).
    """
    try:
        documents = None
        if os.path.exists(PKL_FILENAME):
            with open(PKL_FILENAME, "rb") as file:
                documents = pickle.load(file)

        if documents is None:
            logging.info("Reading documents from directory...")
            documents = read_documents_from_directory(
                DOCUMENTS_DIR,
                TEXT_OUTPUT_DIR,
                TEXT_EXPANDED_DIR,
                checkpoint_callback=persist_documents_snapshot,
            )
        else:
            logging.info("Appending new documents to the existing list...")
            documents = read_documents_from_directory(
                DOCUMENTS_DIR,
                TEXT_OUTPUT_DIR,
                TEXT_EXPANDED_DIR,
                documents,
                checkpoint_callback=persist_documents_snapshot,
            )

            # Debug: Check documents before keyword extraction
            logging.debug(
                f"Number of documents before keyword extraction: {len(documents)}"
            )
            for i, doc in enumerate(documents[:5]):
                logging.debug(f"Document {i} text length: {len(doc['text'])}")

        # Apply keyword extraction
        extract_keywords(documents)

        # Download NLTK data
        download_nltk_data()

        # Save the updated list
        persist_documents_snapshot(documents)

        logging.info(f"Total documents: {len(documents)}")
    except Exception as e:
        logging.exception("An error occurred in main:")
        logging.info(f"An error occurred: {e}")
        sys.exit(1)


if __name__ == "__main__":
    logging.getLogger().setLevel(getattr(logging, LOG_LEVEL, logging.INFO))
    logging.info("Starting document processing script")
    main()
    logging.info("Document processing script completed")
