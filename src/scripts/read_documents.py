import os
import pickle
import pdfplumber
import spacy
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import csv
import re
from docx import Document
from spellchecker import SpellChecker
import wordninja
from sklearn.feature_extraction.text import TfidfVectorizer
import PyPDF2
import logging
# Load spaCy's English model
nlp = spacy.load('en_core_web_sm')
nlp.max_length = 2000000  # or any other suitable value
# Download required NLTK data
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')
# Initialize spellchecker
spell = SpellChecker()
# Suppress specific warnings
import warnings
warnings.filterwarnings("ignore", message="usetex mode requires TeX.")
# Global stopwords
STOP_WORDS = set(stopwords.words('english'))
import docx


# Configure logging
logging.basicConfig(level=logging.INFO, filename='read_documents.log', format='%(asctime)s - %(levelname)s - %(message)s')

# Define base directory
BASE_DIR = r'C:\Users\Aspire5 15 i7 4G2050\ProjectRAG\AviationRAG'

# Define paths
DOCUMENTS_DIR = os.path.join(BASE_DIR, 'data', 'documents')
TEXT_OUTPUT_DIR = os.path.join(BASE_DIR, 'data', 'processed', 'ProcessedText')
TEXT_EXPANDED_DIR = os.path.join(BASE_DIR, 'data', 'processed', 'ProcessedTextExpanded')
PKL_FILENAME = os.path.join(BASE_DIR, 'data', 'raw', 'aviation_corpus.pkl')

# Ensure directories exist
for directory in [TEXT_OUTPUT_DIR, TEXT_EXPANDED_DIR, os.path.dirname(PKL_FILENAME)]:
    if not os.path.exists(directory):
        os.makedirs(directory)
        logging.info(f"Created directory: {directory}")
    else:
        logging.info(f"Directory already exists: {directory}")


# Create custom pipeline component for aviation NER
@spacy.Language.component("aviation_ner")
def aviation_ner(doc):
    logging.info(f"Starting aviation_ner for document: {doc[:50]}...")
    patterns = [
        ("AIRCRAFT_MODEL", r"\b[A-Z]-?[0-9]{1,4}\b"),
        ("AIRPORT_CODE", r"\b[A-Z]{3}\b"),
        ("FLIGHT_NUMBER", r"\b[A-Z]{2,3}\s?[0-9]{1,4}\b"),
        ("AIRLINE", r"\b(American Airlines|Delta Air Lines|United Airlines|Southwest Airlines|Air France|Lufthansa|British Airways)\b"),
        ("AVIATION_ORG", r"\b(FAA|EASA|ICAO|IATA)\b"),
    ]
    
    new_ents = []
    for ent_type, pattern in patterns:
        for match in re.finditer(pattern, doc.text):
            start, end = match.span()
            span = doc.char_span(start, end, label=ent_type)
            if span is not None:
                # Check for overlap with existing entities
                if not any(span.start < ent.end and span.end > ent.start for ent in list(doc.ents) + new_ents):
                    new_ents.append(span)
                    logging.debug(f"Added new entity: {span.text} ({ent_type})")
                    
    doc.ents = list(doc.ents) + new_ents
    logging.info(f"Finished aviation_ner. Added {len(new_ents)} new entities.")
    return doc

# Add the custom component to the pipeline
nlp.add_pipe("aviation_ner", after="ner")

def load_abbreviation_dict():
    abbreviation_dict = {}
    try:
        with open('abbreviations.csv', mode='r') as infile:
            reader = csv.reader(infile)
            for rows in reader:
                if len(rows) < 2:
                    continue
                abbreviation_dict[rows[0].strip()] = rows[1].strip()
    except FileNotFoundError:
        print("Error: The file 'abbreviations.csv' was not found.")
    except Exception as e:
        print(f"An error occurred while loading the abbreviation dictionary: {e}")
    return abbreviation_dict

def split_connected_words_improved(text):
    words = re.findall(r'\w+|\W+', text)
    split_words = []
    for word in words:
        if len(word) > 15 and word.isalnum():
            split_parts = re.findall('[A-Z][a-z]*|[a-z]+|[0-9]+', word)
            split_words.extend(split_parts)
        else:
            split_words.append(word)
    split_words = ' '.join(split_words)
    split_words = ' '.join(wordninja.split(split_words))
    return split_words

def filter_non_sense_strings(text):
    words = text.split()
    cleaned_words = []
    for word in words:
        if re.match(r'^[a-zA-Z]+$', word) and len(set(word.lower())) > 3:
            cleaned_words.append(word)
    return ' '.join(cleaned_words)

def preprocess_text_with_sentences(text):
    doc = nlp(text)
    sentences = []
    for sent in doc.sents:
        cleaned_sentence = ' '.join(
            token.lemma_.lower() for token in sent
            if token.is_alpha and token.text.lower() not in STOP_WORDS
        )
        if cleaned_sentence:
            sentences.append(cleaned_sentence)
    return ' '.join(sentences)

def extract_personal_names(text):
    doc = nlp(text)
    return [ent.text for ent in doc.ents if ent.label_ == 'PERSON']

def extract_entities_and_pos_tags(text):
    doc = nlp(text)
    entities = [(ent.text, ent.label_) for ent in doc.ents]
    pos_tags = [(token.text, token.pos_) for token in doc]
    return entities, pos_tags

def expand_abbreviations_in_text(text, abbreviation_dict):
    words = text.split()
    expanded_words = []
    for word in words:
        if word.lower() in abbreviation_dict:
            expanded_words.append(abbreviation_dict[word.lower()])
        else:
            expanded_words.append(word)
    return ' '.join(expanded_words)

def extract_text_from_pdf_with_pdfplumber(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ''.join([page.extract_text() + '\n' for page in pdf.pages])
            return text
    except Exception as e:
        print(f"Failed to process PDF {pdf_path}: {e}")
        return ""

def extract_keywords(documents, top_n=10):
    texts = [doc['text'] for doc in documents]
    
    # Debug: Check if we have any texts
    logging.debug(f"Number of documents: {len(texts)}")
    if len(texts) == 0:
        logging.error("No texts found in documents")
        return

    # Debug: Check the content of the first few texts
    for i, text in enumerate(texts[:5]):
        logging.debug(f"Text {i} (first 100 chars): {text[:100]}")
        logging.debug(f"Text {i} length: {len(text)}")

    vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
    
    try:
        tfidf_matrix = vectorizer.fit_transform(texts)
    except ValueError as e:
        logging.error(f"ValueError in TfidfVectorizer: {e}")
        logging.error("Vocabulary: " + str(vectorizer.vocabulary_))
        return

def extract_metadata(file_path):
    """
    Extract metadata from PDF or DOCX files.
    
    Args:
        file_path (str): Path to the file.

    Returns:
        dict: Metadata extracted from the file.
    """
    metadata = {
        'title': '',
        'author(s)': '',
        'category': '',
        'tags': ''
    }
    
    # Handle PDF files
    if file_path.endswith('.pdf'):
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            pdf_metadata = reader.metadata or {}
            metadata['title'] = pdf_metadata.get('/Title', '')
            metadata['author(s)'] = pdf_metadata.get('/Author', '')
    
    elif file_path.endswith('.docx'):
        try:
            doc = docx.Document(file_path)
            # Look at the first few paragraphs for metadata
            text = '\n'.join([p.text for p in doc.paragraphs[:5]])
            
            # Use regex to find metadata
            title_match = re.search(r'Title:\s*(.*)', text)
            if title_match:
                metadata['title'] = title_match.group(1).strip()
            
            author_match = re.search(r'Author\(s\):\s*(.*)', text)
            if author_match:
                metadata['author(s)'] = author_match.group(1).strip()
            
            category_match = re.search(r'Category:\s*(.*)', text)
            if category_match:
                metadata['category'] = category_match.group(1).strip()
            
            tags_match = re.search(r'Tags:\s*(.*)', text)
            if tags_match:
                metadata['tags'] = tags_match.group(1).strip()
            
        except Exception as e:
            print(f"Error reading DOCX metadata: {e}")
    
    return metadata

def classify_document(text):
    keywords = {
        'safety': ['safety', 'hazard', 'risk', 'incident', 'accident','system','hazard','emergency'],
        'maintenance': ['maintenance', 'repair', 'overhaul', 'inspection'],
        'operations': ['flight', 'takeoff', 'landing', 'crew', 'pilot','aircraft', 'airplane'],
        'regulations': ['regulation', 'compliance', 'standard', 'rule', 'law'],
        'quality': ['quality', 'performance', 'service', 'customer', 'satisfaction','design'],
    }
    
    text_lower = text.lower()
    scores = {category: sum(1 for word in words if word in text_lower) for category, words in keywords.items()}
    return max(scores, key=scores.get)

def read_documents_from_directory(directory_path, text_output_dir=None, text_expanded_dir=None, existing_documents=None):
    logging.info(f"Starting to read documents from {directory_path}")
    if existing_documents is None:
        existing_documents = []
    
    existing_files = {doc['filename'] for doc in existing_documents}
    new_documents = []
    abbreviation_dict = load_abbreviation_dict()
    lemmatizer = WordNetLemmatizer()

    for filename in os.listdir(directory_path):
        logging.info(f"Processing file: {filename}")
        if filename in existing_files:
            continue

        file_path = os.path.join(directory_path, filename)
        text = ''
        if filename.endswith(".pdf"):
            logging.info(f"Extracting text from PDF: {filename}")
            text = extract_text_from_pdf_with_pdfplumber(file_path)
        elif filename.endswith(".docx"):
            logging.info(f"Extracting text from DOCX: {filename}")
            try:
                doc = Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as e:
                logging.error(f"Failed to process DOCX {filename}: {e}")
                print(f"Failed to process DOCX {filename}: {e}")
                continue
        else:
            logging.warning(f"Skipping unsupported file type: {filename}")
            continue

        logging.info(f"Preprocessing text from {filename}")
        if not text:
            logging.warning(f"No text extracted from {filename}")
            continue

        expanded_text = expand_abbreviations_in_text(text, abbreviation_dict)
        raw_text = expanded_text
        expanded_text = split_connected_words_improved(expanded_text)
        expanded_text = filter_non_sense_strings(expanded_text)
        preprocessed_text = preprocess_text_with_sentences(expanded_text)
        personal_names = extract_personal_names(preprocessed_text)
        entities, pos_tags = extract_entities_and_pos_tags(preprocessed_text)
        tokens = word_tokenize(preprocessed_text)
        cleaned_tokens = [token.lower() for token in tokens if token.isalpha() and len(token) > 2]
        tokens_without_stopwords = [token for token in cleaned_tokens if token not in STOP_WORDS]
        lemmatized_tokens = [lemmatizer.lemmatize(token) for token in tokens_without_stopwords]

        if text_expanded_dir:
            # Remove .docx extension before adding .txt
            clean_filename = filename.replace('.docx', '')
            output_file_path = os.path.join(text_expanded_dir, f'{clean_filename}.txt')
            with open(output_file_path, 'w', encoding='utf-8') as out_file:
                out_file.write(raw_text)
            logging.info(f"Expanded text saved to: {output_file_path}")
            print(f"Expanded text saved to: {output_file_path}")
               
        if text_output_dir:
            # Remove .docx extension before adding .txt
            clean_filename = filename.replace('.docx', '')
            output_file_path = os.path.join(text_output_dir, f'{clean_filename}.txt')
            logging.info(f"Processed text saved to: {output_file_path}")
            with open(output_file_path, 'w', encoding='utf-8') as out_file:
                out_file.write(preprocessed_text)
            print(f"Text saved to: {output_file_path}")

        logging.info(f"Finished processing all documents in {directory_path}")
        metadata = extract_metadata(file_path)
        document_category = classify_document(preprocessed_text)

        # Remove 'category' field if it already exists
        if 'category' in metadata:
            del metadata['category']

        new_documents.append({
            'filename': filename,
            'text': preprocessed_text,
            'tokens': lemmatized_tokens,
            'personal_names': personal_names,
            'entities': entities,
            'pos_tags': pos_tags,
            'metadata': metadata,
            'category': document_category
        })

    return existing_documents + new_documents

def update_existing_documents(documents):
    for doc in documents:
        if 'metadata' not in doc:
            doc['metadata'] = extract_metadata(os.path.join(BASE_DIR, doc['filename']))
        if 'category' not in doc:
            doc['category'] = classify_document(doc['text'])
    return documents

def main():
    try:
        documents = None
        if os.path.exists(PKL_FILENAME):
            with open(PKL_FILENAME, 'rb') as file:
                documents = pickle.load(file)
            documents = update_existing_documents(documents)

        if documents is None:
            print("Reading documents from directory...")
            documents = read_documents_from_directory(DOCUMENTS_DIR, TEXT_OUTPUT_DIR, TEXT_EXPANDED_DIR)
        else:
            print("Appending new documents to the existing list...")
            documents = read_documents_from_directory(DOCUMENTS_DIR, TEXT_OUTPUT_DIR, TEXT_EXPANDED_DIR, documents)

        # Debug: Check documents before keyword extraction
            logging.debug(f"Number of documents before keyword extraction: {len(documents)}")
            for i, doc in enumerate(documents[:5]):
                logging.debug(f"Document {i} text length: {len(doc['text'])}")


        # Apply keyword extraction
        extract_keywords(documents)

        # Save the updated list
        with open(PKL_FILENAME, 'wb') as file:
            pickle.dump(documents, file)

        print(f"Total documents: {len(documents)}")
    except Exception as e:
        logging.exception("An error occurred in main:")
        print(f"An error occurred: {e}")

if __name__ == '__main__':
    logging.basicConfig(level=logging.DEBUG)
    logging.info("Starting document processing script")
    main()
    logging.info("Document processing script completed")

